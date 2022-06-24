from tkinter import ttk
from tkinter.scrolledtext import ScrolledText
from tkinter.ttk import Combobox
from tkinter import filedialog as fd

import Crypto.Util.number as CryNum 
import random
import hashlib #hàm băm md5
import docx
import io
import PyPDF2



from tkinter import * #for gui

def gcd(a, b):
    a = abs(a)
    b = abs(b)
    if a < b:
        a, b = b, a

    while b != 0:
        a, b = b, a % b

    return a


def genRandE(phi):
    e = random.randint(1, phi)
    g = gcd(e, phi)
    # checking if e and )(n) is coprime
    while g != 1:
        e = random.randint(1, phi)
        g = gcd(e, phi)

    return e

def exteuclid(a, b):
    x, old_x = 0, 1
    y, old_y = 1, 0

    while (b != 0):
        quotient = a // b
        a, b = b, a - quotient * b
        old_x, x = x, old_x - quotient * x
        old_y, y = y, old_y - quotient * y
    return old_x

def RSA_genKeys(k):
    p, q = CryNum.getPrime(k), CryNum.getPrime(k)
    while p == q:
        p, q = CryNum.getPrime(k), CryNum.getPrime(k)

    n = p * q
    phi = (p - 1) * (q - 1)
    e = genRandE(phi)
    d = exteuclid(e, phi)
    if d < 0:
        d += phi
    m=n
    return e,n,d,m

def ShowKeys(k):
    new_e, new_n, new_d, new_m = RSA_genKeys(k)
    e.set(new_e)
    n.set(new_n)
    d.set(new_d)
    m.set(new_m)
    e_input.delete(1.0, END)
    e_input.insert(1.0, e.get())
    n_input.delete(1.0, END)
    n_input.insert(1.0, n.get())
    d_input.delete(1.0, END)
    d_input.insert(1.0, d.get())
    m_input.delete(1.0, END)
    m_input.insert(1.0, m.get())

def setKeys():
    keys=open('Keys.txt','w')
    keys.write(
        str(e_input.get(1.0,END)).strip() + "\n" +
        str(n_input.get(1.0,END)).strip() + "\n" +
        str(d_input.get(1.0,END)).strip() + "\n" +
        str(m_input.get(1.0,END)).strip()
        )

def sign():
    new_msg=msg_input.get(1.0,END).strip()
    keysfile=open('Keys.txt')
    keys=keysfile.read()
    keys=keys.split("\n")
    d=int(keys[2])
    m=int(keys[3])
    digest = hashlib.md5(new_msg.encode('utf-8')).hexdigest()
    M = int(digest, 16)
    encDigest = pow(M, d, m)
    sig_input.delete(1.0,END)
    sig_input.insert(1.0,encDigest)
    sig.set(encDigest)
    keysfile.close()

def saveData():
    text_file = fd.askopenfilename(initialdir="C:/Users/DangHoc/Desktop/project/RSA", title="Open Text file", filetypes=(("Text Files", "*.txt"),("Microsoft Word Document","*.docx"),("PDF Files", "*.pdf"),))
    types = text_file.split('.', 1)
    type = types[1]
    if(type == "docx"):
        text_file = docx.Document(text_file)
        message= ""
        for i in range(len(text_file.paragraphs)):
            message += text_file.paragraphs[i].text
            message += '\n'

    if(type == "txt"):
        text_file = io.open(text_file, 'r', encoding='utf8')        
        data = text_file.read()
        data=data.split("\n")
        message = ""
        for i in range(len(data)-1):
            message+=data[i]
            if i!=len(data)-1:
                message+='\n'
        text_file.close()
    
    if(type == "pdf"):
        text_file = PyPDF2.PdfFileReader(text_file)
        for i in range(len(text_file.pages)):
            message += text_file.getPage(i).extractText()
    msg_input.delete(1.0,END)
    sig_input.delete(1.0,END)
    msg_input.insert(1.0, message)

def readKeys():
    keysfile=open('Keys.txt')
    keys=keysfile.read()
    keys=keys.split("\n")
    e.set(keys[0])
    n.set(keys[1])
    d.set(keys[2])
    m.set(keys[3])
    e_input.delete(1.0, END)
    e_input.insert(1.0, e.get())
    n_input.delete(1.0, END)
    n_input.insert(1.0, n.get())
    d_input.delete(1.0, END)
    d_input.insert(1.0, d.get())
    m_input.delete(1.0, END)
    m_input.insert(1.0, m.get())
    keysfile.close()


def readData():
    text_file = fd.askopenfilename(initialdir="C:/Users/DangHoc/Desktop/project/RSA", title="Open Text file", filetypes=(("Text Files", "*.txt"),("Microsoft Word Document","*.docx"),("PDF Files", "*.pdf"),))
    types = text_file.split('.', 1)
    type = types[1]
    if(type == "docx"):
        text_file = docx.Document(text_file)
        message= ""
        for i in range(len(text_file.paragraphs)):
            message += text_file.paragraphs[i].text
            message += '\n'

    if(type == "txt"):
        text_file = io.open(text_file, 'r', encoding='utf8')        
        data = text_file.read()
        data=data.split("\n")
        message = ""
        for i in range(len(data)-1):
            message+=data[i]
            if i!=len(data)-1:
                message+='\n'
        text_file.close()

    if(type == "pdf"):
        text_file = PyPDF2.PdfFileReader(text_file)
        for i in range(len(text_file.pages)):
            message += text_file.getPage(i).extractText()

    msg2_input.delete(1.0,END)
    sig2_input.delete(1.0,END)
    msg2_input.insert(1.0,message)
    # sig2_input.insert(1.0,int(data[-1]))
    res.set("Click button để xác thực.")
    

def verify():
    new_msg=msg2.get()
    file=open('Keys.txt')
    keys=file.read()
    file.close()
    keys=keys.split("\n")
    e=int(keys[0])
    n=int(keys[3])
    M = int(hashlib.md5(msg2_input.get(1.0,END).strip().encode('utf-8')).hexdigest(),16)%n
    digest=sig2_input.get(1.0,END).strip()
    o = pow(int(digest), e, n)
    # print(M,o)
    if(M==o):
        # print(1)
        res.set("Chữ kí xác thực không có gì thay đổi!")
    else:
        # print(2)
        res.set("Chữ kí không khớp!!")

    file.close()


root = Tk()
root.title("Chữ kí số RSA")
root.geometry("920x750")
# root.configure(bg='white')
s = ttk.Style(root)
s.configure("TNotebook", tabposition='n',background='#6C6C6C')
s.configure("TFrame",background='#FAEBD7')
notebook = ttk.Notebook(root,padding=20);
notebook.pack(expand=1, fill="both")
frame1 = ttk.Frame(notebook,relief=GROOVE,padding=5)
frame2 = ttk.Frame(notebook,relief=GROOVE,padding=5)
frame3 = ttk.Frame(notebook,relief=GROOVE,padding=5)
frame4 = ttk.Frame(notebook,relief=GROOVE,padding=5)
notebook.add(frame1, text='   Tạo khóa   ')
notebook.add(frame2, text='   Kí văn bản   ')
notebook.add(frame3, text='   Xác nhận chữ kí   ')
notebook.add(frame4, text='   Cài đặt   ')
keysizes=[128,256,512,1024]
current_keysize=IntVar();
current_keysize.set(1024);
current_Algorithm=StringVar();
current_Algorithm.set("RSA");
Generate_Keys_Label=Label(frame1,text="  Tạo khóa : ",font=('Arial', 14, 'bold'),background='#FAEBD7')
Generate_Keys_Label.grid(padx=5,pady=10,row=0,column=0)
Public_Key_Label = Label(frame1,text="  Public Key: ",font=('Arial', 12),background='#FAEBD7')
Public_Key_Label.grid(padx=5,pady=10,row=2, column=0)
Label(frame1,text="e: ",font=('Arial', 12),background='#FAEBD7').grid(padx=5,pady=10,row=3,column=0)
e=IntVar()
e_input=ScrolledText(frame1,width=80,height=5)
e_input.grid(padx=5,pady=10,row=3, column=1,sticky="E",columnspan=4)
Label(frame1,text="n: ",font=('Arial', 12),background='#FAEBD7').grid(padx=5,pady=10,row=4,column=0)
n=IntVar()
n_input=ScrolledText(frame1,width=80,height=5)
n_input.grid(padx=5,pady=10,row=4, column=1,sticky="E",columnspan=4)
Private_Key_Label=Label(frame1,text="  Private Key: ",font=('Arial', 12),background='#FAEBD7')
Private_Key_Label.grid(padx=5,pady=10,row=5,column=0)
Label(frame1,text="d: ",font=('Arial', 12),background='#FAEBD7').grid(padx=5,pady=10,row=6,column=0)
d=IntVar()
d_input=ScrolledText(frame1,width=80,height=5)
d_input.grid(padx=5,pady=10,row=6, column=1,sticky="E",columnspan=4)
m=IntVar()
Label(frame1,text="m: ",font=('Arial', 12),background='#FAEBD7').grid(padx=5,pady=10,row=7,column=0)
m_input=ScrolledText(frame1,width=80,height=5)
m_input.grid(padx=5,pady=10,row=7, column=1,sticky="E",columnspan=4)
#Label(frame1,text=" ").grid(padx=5,pady=10,row=8,column=0)
Read_Keys_Button=Button(frame1,text = " Đọc Keys ",font=('Arial', 12),activebackground='Coral',background='Coral',relief=FLAT,command=readKeys)
Read_Keys_Button.grid(ipadx=10,ipady=10,padx=5,pady=10,row=9,column=1)
Generate_Keys_Button=Button(frame1,text = " Tạo khóa ",font=('Arial', 12),activebackground='Coral',background='LightSeaGreen',relief=FLAT,command=lambda: ShowKeys(int(current_keysize.get()//2)))
Generate_Keys_Button.grid(ipadx=10,ipady=10,padx=5,pady=10,row=9,column=2)
Set_Keys_Button=Button(frame1,text = " Set Keys ",font=('Arial', 12),activebackground='Coral',background='LightSkyBlue',relief=FLAT,command=setKeys)
Set_Keys_Button.grid(ipadx=10,ipady=10,padx=5,pady=10,row=9,column=3)
#Label(frame1,text=" ").grid(padx=5,pady=10,row=10,column=0)
Label(frame2,text="  Kí:               ",font=('Arial', 14, 'bold'),background='#FAEBD7').grid(padx=5,pady=10,row=11,column=0)
msg=StringVar(frame2,"")
#Label(frame2,text=" ").grid(padx=5,pady=10,row=12,column=0)
Label(frame2,text=" Văn bản cần kí:    ",font=('Arial', 12),background='#FAEBD7').grid(padx=5,pady=10,row=13,column=0)
msg_input=ScrolledText(frame2,width=80,height=10)
msg_input.grid(padx=5,pady=10,row=13, column=1,sticky="E",columnspan=4)
Sign_Button=Button(frame2,text = " Kí ",font=('Arial', 12),activebackground='Coral',background='LightSeaGreen',relief=FLAT,command=sign)
Sign_Button.grid(ipadx=10,ipady=10,padx=5,pady=10,row=14,column=2)
sig=IntVar()
Label(frame2,text=" Chữ kí:      ",font=('Arial', 12),background='#FAEBD7').grid(padx=5,pady=10,row=15,column=0)
sig_input=ScrolledText(frame2,width=80,height=5)
sig_input.grid(padx=5,pady=10,row=15, column=1,sticky="E",columnspan=4)
Save_Data_Button=Button(frame2,text = " Đọc Data ",font=('Arial', 12),activebackground='Coral',background='LightSkyBlue',relief=FLAT,command=saveData)
Save_Data_Button.grid(ipadx=10,ipady=10,padx=5,pady=10,row=16,column=2)
#Label(frame2,text=" ").grid(padx=5,pady=10,row=17,column=0)
Label(frame3,text="  Xác thực chữ kí: ",font=('Arial', 14, 'bold'),background='#FAEBD7').grid(padx=5,pady=10,row=14,column=0)
msg2=StringVar()
#Label(frame3,text=" ").grid(padx=5,pady=10,row=19,column=0)
Read_Data_Button=Button(frame3,text = " Đọc Data ",font=('Arial', 12),activebackground='Coral',background='LightSeaGreen',relief=FLAT,command=readData)
Read_Data_Button.grid(ipadx=10,ipady=10,padx=5,pady=10,row=20,column=0,sticky="E")
Label(frame3,text=" Văn bản cần xác nhận: ",font=('Arial', 12),background='#FAEBD7').grid(padx=5,pady=10,row=21,column=0)
msg2_input=ScrolledText(frame3,width=80,height=10)
msg2_input.grid(padx=5,pady=10,row=21, column=1,sticky="E",columnspan=4)
sig2=IntVar()
Label(frame3,text=" Chữ kí: ",font=('Arial', 12),background='#FAEBD7').grid(padx=5,pady=10,row=23,column=0)
sig2_input=ScrolledText(frame3,width=80,height=5)
sig2_input.grid(padx=5,pady=10,row=23, column=1,sticky="E",columnspan=4)
Verify_Button=Button(frame3,text = " Xác thực chữ kí ",activebackground='Coral',font=('Arial', 12),background='LightSkyBlue',relief=FLAT,command=verify)
Verify_Button.grid(ipadx=10,ipady=10,padx=5,pady=10,row=24,column=1)
res=StringVar(frame3,"Click button để xác thực.")
Label(frame3,textvariable=res,font=('Arial', 14),background='#FAEBD7').grid(padx=5,pady=10,row=24,column=2)
Label(frame4,text=" Cài đặt: ",font=('Arial', 14, 'bold'),background='#FAEBD7').grid(padx=5,pady=10,row=0,column=0,sticky='W')
Label(frame4,text=" Chọn độ dài Key(bits>=8): ",font=('Arial', 12),background='#FAEBD7').grid(padx=5,pady=10,row=1,column=0,columnspan=3)
KeySizes_Drop = Combobox( frame4 ,textvariable=current_keysize ,values=keysizes)
#KeySizes_Drop.config(background='Coral')
#KeySizes_Drop['menu'].config(background='LightSkyBlue')
KeySizes_Drop.grid(padx=5,pady=10,row=1,column=3)
Label(frame4,text=" Select Algorithm: ",font=('Arial', 12),background='#FAEBD7').grid(padx=5,pady=10,row=2,column=0,columnspan=3)

Radiobutton(frame4,text=" RSA ",variable=current_Algorithm,value="RSA",indicator=0,activebackground='Coral',background='LightSkyBlue',font=('Arial', 12)).grid(ipadx=10,ipady=10,padx=5,pady=10,row=2,column=4)

root.mainloop()